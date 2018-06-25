#!/usr/bin/python3
# -*- encoding: utf-8 -*-

import docx

doc = docx.Document('demo.docx')
print(doc.paragraphs[1].runs[0].text)
print(doc.paragraphs[1].runs[0].style)
print(doc.paragraphs[1].runs[0].underline)

# スタイルなどを変更する
# `doc.paragraphs[1].runs[0].style = 'Quote' # NG
doc.paragraphs[1].runs[0].underline = True

doc.save('new_demo.docx')

