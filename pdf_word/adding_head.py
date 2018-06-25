#!/usr/bin/python3
# -*- encoding: utf-8 -*-

import docx

doc = docx.Document()
doc.add_heading('Tilteee', 0)
doc.add_heading('Tilteeee', 1)
doc.add_heading('Tilteeeee', 2)
doc.add_heading('Tilteeeeee', 3)
doc.add_paragraph('This is on the first page!')
#doc.paragraphs[0].runs[0].add_break(docx.text.WD_BREAK.PAGE) # NG
doc.add_paragraph('This is on the second page!')
doc.save('headings.docx')

