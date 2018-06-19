#!/usr/bin/python3
# -*- encoding: utf-8 -*-

import docx

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)


doc = docx.Document('demo.docx')
print('number of paragraphs is ' + str(len(doc.paragraphs)) + '.')

# パラグラフ単位でテキスト抽出
print(doc.paragraphs[1].text)

# パラグラフ内のrunオブジェクト単位で抽出(スタイル区切りでテキスト抽出？)
print(doc.paragraphs[1].runs[0].text)
print(doc.paragraphs[1].runs[1].text)
print(doc.paragraphs[1].runs[2].text)
print(doc.paragraphs[1].runs[3].text)

# ファイル内の全テキストを出力する
print('\nll text is below.\n\n')
print(getText('demo.docx'))

